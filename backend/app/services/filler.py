"""
自动填表模块
支持 Excel 和 Word 模板的自动填充
"""
import os
import tempfile
import pandas as pd
from docx import Document
from typing import Dict, Any, List


def fill_template(template_path: str, matches: Dict[str, Dict[str, Any]]) -> str:
    """
    根据匹配结果填充模板
    
    Args:
        template_path: 模板文件路径
        matches: 字段到实体的映射
    
    Returns:
        生成的结果文件路径
    """
    file_ext = os.path.splitext(template_path)[1].lower()
    
    if file_ext == '.xlsx':
        return _fill_excel_template(template_path, matches)
    elif file_ext == '.docx':
        return _fill_word_template(template_path, matches)
    else:
        raise ValueError(f"不支持的模板格式: {file_ext}")


def _fill_excel_template(template_path: str, matches: Dict[str, Dict[str, Any]]) -> str:
    """
    填充 Excel 模板
    """
    # 创建临时文件作为输出
    with tempfile.NamedTemporaryFile(suffix='.xlsx', delete=False) as tmp:
        output_path = tmp.name
    
    try:
        # 读取模板
        df = pd.read_excel(template_path)
        
        # 填充数据
        for field, match_info in matches.items():
            # 查找字段所在的列
            if field in df.columns:
                # 填充第一行
                df[field] = match_info['entity']['value']
        
        # 保存结果
        df.to_excel(output_path, index=False)
        
    except Exception as e:
        print(f"填充 Excel 模板时出错: {e}")
        os.remove(output_path)
        raise
    
    return output_path


def _fill_word_template(template_path: str, matches: Dict[str, Dict[str, Any]]) -> str:
    """
    填充 Word 模板
    """
    # 创建临时文件作为输出
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        output_path = tmp.name
    
    try:
        # 读取模板
        doc = Document(template_path)
        
        # 填充段落
        for paragraph in doc.paragraphs:
            for field, match_info in matches.items():
                if f"{{{field}}}" in paragraph.text:
                    paragraph.text = paragraph.text.replace(f"{{{field}}}", match_info['entity']['value'])
        
        # 填充表格
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for field, match_info in matches.items():
                        if f"{{{field}}}" in cell.text:
                            cell.text = cell.text.replace(f"{{{field}}}", match_info['entity']['value'])
        
        # 保存结果
        doc.save(output_path)
        
    except Exception as e:
        print(f"填充 Word 模板时出错: {e}")
        os.remove(output_path)
        raise
    
    return output_path


def extract_template_fields(template_path: str) -> List[str]:
    """
    从模板中提取字段
    
    Args:
        template_path: 模板文件路径
    
    Returns:
        字段列表
    """
    file_ext = os.path.splitext(template_path)[1].lower()
    
    if file_ext == '.xlsx':
        return _extract_excel_fields(template_path)
    elif file_ext == '.docx':
        return _extract_word_fields(template_path)
    else:
        raise ValueError(f"不支持的模板格式: {file_ext}")


def _extract_excel_fields(template_path: str) -> List[str]:
    """
    从 Excel 模板中提取字段
    """
    try:
        df = pd.read_excel(template_path)
        return list(df.columns)
    except Exception as e:
        print(f"提取 Excel 字段时出错: {e}")
        return []


def _extract_word_fields(template_path: str) -> List[str]:
    """
    从 Word 模板中提取字段
    """
    fields = []
    
    try:
        doc = Document(template_path)
        
        # 从段落中提取字段
        for paragraph in doc.paragraphs:
            text = paragraph.text
            # 查找 {field} 格式的字段
            import re
            matches = re.findall(r'\{([^}]+)\}', text)
            fields.extend(matches)
        
        # 从表格中提取字段
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text
                    matches = re.findall(r'\{([^}]+)\}', text)
                    fields.extend(matches)
        
        # 去重
        fields = list(set(fields))
        
    except Exception as e:
        print(f"提取 Word 字段时出错: {e}")
    
    return fields