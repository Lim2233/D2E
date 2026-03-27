"""
文档解析模块
支持 Word/Excel/Markdown/TXT 转换为统一的 JSON 结构
"""
import os
import json
import markdown
import pandas as pd
from docx import Document
from typing import Dict, List, Any
import uuid


def parse_document(file_path: str) -> Dict[str, Any]:
    """
    解析文档为统一的 JSON 结构
    
    Args:
        file_path: 文件路径
    
    Returns:
        包含 doc_id、paragraphs、tables、raw_text 的字典
    """
    doc_id = str(uuid.uuid4())
    paragraphs = []
    tables = []
    raw_text = ""
    
    try:
        file_ext = os.path.splitext(file_path)[1].lower()
        
        if file_ext == '.docx':
            paragraphs, tables, raw_text = _parse_docx(file_path)
        elif file_ext == '.xlsx':
            paragraphs, tables, raw_text = _parse_xlsx(file_path)
        elif file_ext == '.md':
            paragraphs, tables, raw_text = _parse_md(file_path)
        elif file_ext == '.txt':
            paragraphs, tables, raw_text = _parse_txt(file_path)
        else:
            raise ValueError(f"不支持的文件格式: {file_ext}")
    
    except Exception as e:
        print(f"解析文件时出错: {e}")
        # 返回空结构以避免整个流程失败
        pass
    
    return {
        "doc_id": doc_id,
        "paragraphs": paragraphs,
        "tables": tables,
        "raw_text": raw_text
    }


def _parse_docx(file_path: str) -> tuple:
    """
    解析 Word 文档
    """
    paragraphs = []
    tables = []
    raw_text = ""
    
    try:
        doc = Document(file_path)
        
        # 提取段落
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)
                raw_text += text + "\n"
        
        # 提取表格
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            if table_data:
                tables.append(table_data)
    except Exception as e:
        # 备用方案：使用 zipfile 直接读取 document.xml
        try:
            import zipfile
            from xml.etree import ElementTree as ET
            
            with zipfile.ZipFile(file_path, 'r') as z:
                xml_content = z.read('word/document.xml')
            
            # 解析 XML
            root = ET.fromstring(xml_content)
            
            # Word 命名空间
            ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
            
            # 提取所有段落文本
            for para in root.findall('.//w:p', ns):
                texts = []
                for t in para.findall('.//w:t', ns):
                    if t.text:
                        texts.append(t.text)
                if texts:
                    text = ''.join(texts).strip()
                    if text:
                        paragraphs.append(text)
                        raw_text += text + "\n"
        except Exception as e2:
            print(f"解析 Word 时出错: {e}, 备用方案也失败: {e2}")
    
    return paragraphs, tables, raw_text


def _parse_xlsx(file_path: str) -> tuple:
    """
    解析 Excel 文档
    """
    paragraphs = []
    tables = []
    raw_text = ""
    
    try:
        # 使用pandas读取
        df = pd.read_excel(file_path)
        
        # 将表格转换为三维列表
        table_data = [df.columns.tolist()] + df.values.tolist()
        # 转换所有值为字符串
        table_data = [[str(cell) if cell is not None else "" for cell in row] for row in table_data]
        tables.append(table_data)
        
        # 提取文本内容用于 raw_text
        sheet_text = "Sheet: Sheet1\n"
        for _, row in df.iterrows():
            row_text = " ".join([str(cell) for cell in row if cell is not None])
            if row_text.strip():
                sheet_text += row_text + "\n"
        raw_text += sheet_text + "\n"
    except Exception as e:
        print(f"解析 Excel 时出错: {e}")
    
    return paragraphs, tables, raw_text


def _parse_md(file_path: str) -> tuple:
    """
    解析 Markdown 文档
    """
    paragraphs = []
    tables = []
    raw_text = ""
    
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # 简单处理段落，去掉Markdown格式
        lines = content.split('\n')
        processed_lines = []
        for line in lines:
            # 去掉Markdown格式
            line = line.strip()
            # 去掉标题符号
            line = line.lstrip('#*_`~[]()')
            # 去掉粗体和斜体标记
            line = line.replace('**', '').replace('*', '').replace('_', '')
            if line:
                paragraphs.append(line)
                processed_lines.append(line)
        
        # 生成处理后的raw_text
        raw_text = '\n'.join(processed_lines)
    except Exception as e:
        print(f"解析 Markdown 时出错: {e}")
    
    return paragraphs, tables, raw_text


def _parse_txt(file_path: str) -> tuple:
    """
    解析 TXT 文档
    """
    paragraphs = []
    tables = []
    raw_text = ""
    
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # 提取原始文本
        raw_text = content
        
        # 按换行符分割段落
        lines = content.split('\n')
        for line in lines:
            if line.strip():
                paragraphs.append(line.strip())
    except Exception as e:
        print(f"解析 TXT 时出错: {e}")
    
    return paragraphs, tables, raw_text
