"""
测试文档解析模块
"""
import os
import tempfile
import unittest
from app.services.parser import parse_document


class TestParser(unittest.TestCase):
    """
    测试文档解析模块
    """
    
    def test_parse_txt(self):
        """
        测试解析 TXT 文件
        """
        # 创建临时 TXT 文件
        with tempfile.NamedTemporaryFile(suffix='.txt', delete=False) as tmp:
            tmp.write(b"Hello World\nThis is a test.")
            temp_path = tmp.name
        
        try:
            # 解析文件
            result = parse_document(temp_path)
            
            # 验证结果
            self.assertIsNotNone(result)
            self.assertIn('doc_id', result)
            self.assertIn('paragraphs', result)
            self.assertIn('tables', result)
            self.assertIn('raw_text', result)
            
            # 验证内容
            self.assertEqual(len(result['paragraphs']), 2)
            self.assertEqual(result['paragraphs'][0], "Hello World")
            self.assertEqual(result['paragraphs'][1], "This is a test.")
            self.assertEqual(len(result['tables']), 0)
            self.assertIn("Hello World", result['raw_text'])
            self.assertIn("This is a test.", result['raw_text'])
            
        finally:
            # 清理临时文件
            if os.path.exists(temp_path):
                os.remove(temp_path)
    
    def test_parse_md(self):
        """
        测试解析 Markdown 文件
        """
        # 创建临时 MD 文件
        with tempfile.NamedTemporaryFile(suffix='.md', delete=False) as tmp:
            tmp.write(b"# Test\n\nThis is a **test**.")
            temp_path = tmp.name
        
        try:
            # 解析文件
            result = parse_document(temp_path)
            
            # 验证结果
            self.assertIsNotNone(result)
            self.assertIn('doc_id', result)
            self.assertIn('paragraphs', result)
            self.assertIn('tables', result)
            self.assertIn('raw_text', result)
            
            # 验证内容
            self.assertGreater(len(result['paragraphs']), 0)
            self.assertEqual(len(result['tables']), 0)
            self.assertIn("Test", result['raw_text'])
            self.assertIn("This is a test.", result['raw_text'])
            
        finally:
            # 清理临时文件
            if os.path.exists(temp_path):
                os.remove(temp_path)
    
    def test_parse_docx(self):
        """
        测试解析 Word 文件
        """
        # 创建临时 DOCX 文件
        from docx import Document
        
        doc = Document()
        doc.add_paragraph("Hello Word")
        doc.add_paragraph("This is a test.")
        
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
            temp_path = tmp.name
        
        doc.save(temp_path)
        
        try:
            # 解析文件
            result = parse_document(temp_path)
            
            # 验证结果
            self.assertIsNotNone(result)
            self.assertIn('doc_id', result)
            self.assertIn('paragraphs', result)
            self.assertIn('tables', result)
            self.assertIn('raw_text', result)
            
            # 验证内容
            self.assertGreater(len(result['paragraphs']), 0)
            self.assertEqual(len(result['tables']), 0)
            self.assertIn("Hello Word", result['raw_text'])
            self.assertIn("This is a test.", result['raw_text'])
            
        finally:
            # 清理临时文件
            if os.path.exists(temp_path):
                os.remove(temp_path)
    
    def test_parse_xlsx(self):
        """
        测试解析 Excel 文件
        """
        # 创建临时 XLSX 文件
        import pandas as pd
        
        df = pd.DataFrame({
            'Name': ['Alice', 'Bob'],
            'Age': [25, 30]
        })
        
        with tempfile.NamedTemporaryFile(suffix='.xlsx', delete=False) as tmp:
            temp_path = tmp.name
        
        df.to_excel(temp_path, index=False)
        
        try:
            # 解析文件
            result = parse_document(temp_path)
            
            # 验证结果
            self.assertIsNotNone(result)
            self.assertIn('doc_id', result)
            self.assertIn('paragraphs', result)
            self.assertIn('tables', result)
            self.assertIn('raw_text', result)
            
            # 验证内容
            self.assertEqual(len(result['tables']), 1)
            self.assertIn("Alice", result['raw_text'])
            self.assertIn("Bob", result['raw_text'])
            
        finally:
            # 清理临时文件
            if os.path.exists(temp_path):
                os.remove(temp_path)
    
    def test_parse_invalid_format(self):
        """
        测试解析无效格式文件
        """
        # 创建临时无效格式文件
        with tempfile.NamedTemporaryFile(suffix='.invalid', delete=False) as tmp:
            tmp.write(b"Invalid content")
            temp_path = tmp.name
        
        try:
            # 解析文件
            result = parse_document(temp_path)
            
            # 验证结果（应该返回空结构）
            self.assertIsNotNone(result)
            self.assertIn('doc_id', result)
            self.assertIn('paragraphs', result)
            self.assertIn('tables', result)
            self.assertIn('raw_text', result)
            
        finally:
            # 清理临时文件
            if os.path.exists(temp_path):
                os.remove(temp_path)


if __name__ == '__main__':
    unittest.main()